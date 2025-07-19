# main.py
# Agentic Framework Backend for the AI Business Analyst - FINAL

from flask import Flask, request, jsonify, redirect, send_file
from flask_cors import CORS
import os
from io import BytesIO
import requests
import json
import uuid
import base64
import io
import re
from datetime import datetime

# --- File Parsing Libraries ---
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import PyPDF2

# --- Azure SDKs ---
from azure.communication.email import EmailClient

# --- Database imports ---
from database import (
    init_db, get_db, save_document_to_db, save_analysis_to_db,
    add_to_vector_db, search_vector_db, Document, Analysis
)

# --- Configuration imports ---
from config import (
    GEMINI_API_URL, ACS_CONNECTION_STRING, ACS_SENDER_ADDRESS,
    APPROVAL_RECIPIENT_EMAIL, BACKEND_BASE_URL, ADO_ORGANIZATION_URL,
    ADO_PROJECT_NAME, ADO_PAT
)

# Initialize the Flask application
app = Flask(__name__)

# Initialize database
init_db()

# --- Configuration ---
CORS(app, resources={r"/api/*": {"origins": "*"}})

# --- In-memory storage (for demo purposes) ---
approval_statuses = {}
documents_storage = []
past_analyses_storage = []

# --- Document Management ---
@app.route("/api/upload_document", methods=['POST'])
def upload_document():
    """Upload and store a document"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not file.filename.lower().endswith(('.pdf', '.docx')):
            return jsonify({"error": "Only PDF and DOCX files are allowed"}), 400

        # Extract content from file
        file_content, images, error = agent_extract_content(io.BytesIO(file.read()), file.filename)
        if error:
            return jsonify({"error": f"Failed to extract content: {error}"}), 500

        # Create document record
        doc_id = str(uuid.uuid4())
        document_data = {
            "id": doc_id,
            "name": file.filename,
            "uploadDate": datetime.now().isoformat(),
            "fileType": file.filename.split('.')[-1].lower(),
            "size": len(file.read()),
            "status": "uploaded"
        }
        
        # Reset file pointer
        file.seek(0)
        
        # Save to database
        db = next(get_db())
        try:
            save_document_to_db(db, document_data, f"uploads/{doc_id}_{file.filename}", file_content)
            
            # Add to vector database
            add_to_vector_db(
                content=file_content,
                metadata={
                    "id": doc_id,
                    "name": file.filename,
                    "type": "document",
                    "upload_date": document_data["uploadDate"]
                },
                collection_name="documents"
            )
            
        finally:
            db.close()
        
        return jsonify(document_data), 201
        
    except Exception as e:
        print(f"Error uploading document: {e}")
        return jsonify({"error": f"Failed to upload document: {str(e)}"}), 500

@app.route("/api/documents", methods=['GET'])
def get_documents():
    """Get all uploaded documents"""
    db = next(get_db())
    try:
        documents = db.query(Document).all()
        return jsonify([{
            "id": doc.id,
            "name": doc.name,
            "uploadDate": doc.upload_date.isoformat(),
            "fileType": doc.file_type,
            "status": doc.status
        } for doc in documents])
    finally:
        db.close()

@app.route("/api/documents/<doc_id>", methods=['GET'])
def get_document(doc_id):
    """Get a specific document"""
    db = next(get_db())
    try:
        document = db.query(Document).filter(Document.id == doc_id).first()
        if not document:
            return jsonify({"error": "Document not found"}), 404
        return jsonify({
            "id": document.id,
            "name": document.name,
            "uploadDate": document.upload_date.isoformat(),
            "fileType": document.file_type,
            "content": document.content,
            "status": document.status
        })
    finally:
        db.close()

# --- Past Analyses Management ---
@app.route("/api/analyses", methods=['GET'])
def get_past_analyses():
    """Get all past analyses"""
    db = next(get_db())
    try:
        analyses = db.query(Analysis).all()
        return jsonify([{
            "id": analysis.id,
            "title": analysis.title,
            "date": analysis.date.isoformat(),
            "status": analysis.status,
            "originalText": analysis.original_text
        } for analysis in analyses])
    finally:
        db.close()

@app.route("/api/analyses/<analysis_id>", methods=['GET'])
def get_analysis(analysis_id):
    """Get a specific analysis"""
    db = next(get_db())
    try:
        analysis = db.query(Analysis).filter(Analysis.id == analysis_id).first()
        if not analysis:
            return jsonify({"error": "Analysis not found"}), 404
        return jsonify({
            "id": analysis.id,
            "title": analysis.title,
            "date": analysis.date.isoformat(),
            "status": analysis.status,
            "originalText": analysis.original_text,
            "results": analysis.results
        })
    finally:
        db.close()

def save_analysis_results(results, original_text, filename, user_email=None):
    """Save analysis results to database"""
    analysis_id = str(uuid.uuid4())
    analysis_data = {
        "id": analysis_id,
        "title": f"Analysis of {filename}" if filename else "Text Analysis",
        "date": datetime.now().isoformat(),
        "status": "completed",
        "originalText": original_text[:500] + "..." if len(original_text) > 500 else original_text,
        "results": results,
        "user_email": user_email
    }
    
    # Save to database
    db = next(get_db())
    try:
        save_analysis_to_db(db, analysis_data)
        
        # Add to vector database
        add_to_vector_db(
            content=original_text,
            metadata={
                "id": analysis_id,
                "title": analysis_data["title"],
                "type": "analysis",
                "date": analysis_data["date"]
            },
            collection_name="analyses"
        )
        
    finally:
        db.close()
    
    return analysis_id

# --- Helper function to convert markdown to DOCX ---
def markdown_to_docx(markdown_content):
    """Convert markdown content to DOCX format"""
    doc = Document()
    
    # Split content into lines
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level == 1:
                doc.add_heading(text, 0)
            elif level == 2:
                doc.add_heading(text, 1)
            elif level == 3:
                doc.add_heading(text, 2)
            elif level == 4:
                doc.add_heading(text, 3)
            else:
                doc.add_heading(text, 4)
        # Handle bold text
        elif '**' in line:
            # Simple bold handling - replace **text** with bold
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold text
                    p.add_run(part).bold = True
                else:  # Regular text
                    p.add_run(part)
        # Handle lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        # Handle code blocks
        elif line.startswith('```'):
            continue  # Skip code block markers
        # Handle regular paragraphs
        else:
            doc.add_paragraph(line)
    
    return doc

# --- File Extraction Agent ---

def agent_extract_content(file_stream, filename):
    """
    This agent is responsible for extracting all text and images from a document.
    """
    print("AGENT [Extractor]: Starting content extraction...")
    if filename.endswith('.docx'):
        try:
            document = Document(file_stream)
            text_content = "\n".join([para.text for para in document.paragraphs])
            images = []
            for rel in document.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    base64_image = base64.b64encode(image_bytes).decode('utf-8')
                    images.append({"mime_type": image_part.content_type, "data": base64_image})
            print(f"AGENT [Extractor]: Extracted text and {len(images)} images from DOCX.")
            return text_content, images, None
        except Exception as e:
            error = f"Error reading docx file: {e}"
            print(f"AGENT [Extractor]: {error}")
            return None, None, error
            
    elif filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text = "".join(page.extract_text() for page in pdf_reader.pages if page.extract_text())
            print("AGENT [Extractor]: Extracted text from PDF. Image extraction from PDF is not supported in this version.")
            return text, [], None
        except Exception as e:
            error = f"Error reading pdf file: {e}"
            print(f"AGENT [Extractor]: {error}")
            return None, None, error
            
    return None, None, "Unsupported file type"

# --- Gemini API Agent Caller ---

def call_generative_agent(prompt_parts, is_json=False):
    """
    A generic function to call the Gemini API. This acts as the core for all specialized agents.
    """
    headers = {'Content-Type': 'application/json'}
    payload = {"contents": [{"parts": prompt_parts}]}
    if is_json:
        payload["generationConfig"] = {"responseMimeType": "application/json"}

    try:
        response = requests.post(GEMINI_API_URL, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        result = response.json()
        if (result.get('candidates') and result['candidates'][0].get('content') and result['candidates'][0]['content'].get('parts')):
            return result['candidates'][0]['content']['parts'][0]['text'], None
        else:
            return None, f"API response missing expected content: {result}"
    except Exception as e:
        return None, f"An unexpected error occurred contacting the AI service: {e}"

# --- Specialized Generative Agents ---

def agent_planner(text_content, images):
    print("AGENT [Planner]: Starting analysis to create a high-level plan...")
    prompt_text = f"""Analyze the following business requirements document, including its text and any images. Create a concise, high-level plan that summarizes the key components, user roles, and primary user flows. This plan will be used by other specialist agents to generate detailed documents.

--- DOCUMENT TEXT ---
{text_content}
--- DOCUMENT IMAGES ---
"""
    prompt_parts = [{"text": prompt_text}]
    for img in images:
        prompt_parts.append({"inline_data": img})
    
    plan, error = call_generative_agent(prompt_parts)
    if error:
        print(f"AGENT [Planner]: FAILED. {error}")
    else:
        print("AGENT [Planner]: Plan created successfully.")
    return plan, error

def agent_trd_writer(plan, original_text):
    print("AGENT [TRD Writer]: Generating Technical Requirements Document...")
    prompt = f"""Using the following high-level plan and the original requirements text, write a comprehensive Technical Requirements Document (TRD) in Markdown format. Ensure it is detailed and well-structured.

--- HIGH-LEVEL PLAN ---
{plan}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    return call_generative_agent([{"text": prompt}])

def agent_diagrammer(plan, diagram_type):
    print(f"AGENT [Diagrammer]: Generating {diagram_type} diagram...")
    syntax = "graph TD" if diagram_type == "HLD" else "sequenceDiagram"
    prompt = f"""Based on the following high-level plan, generate a {diagram_type} diagram. The output must be ONLY the Mermaid code block for a `{syntax}` diagram. Do not include any explanatory text or markdown backticks.

--- HIGH-LEVEL PLAN ---
{plan}
"""
    return call_generative_agent([{"text": prompt}])

def agent_backlog_creator(plan, original_text):
    print("AGENT [Backlog Creator]: Generating project backlog...")
    prompt = f"""Based on the following high-level plan and original requirements, generate a hierarchical project backlog. 
The output MUST be a single, well-formed JSON object. 
The JSON object must have a single top-level key named "backlog".
The value of "backlog" MUST be a JSON array of Epic objects.

Example of the required output structure:
{{
  "backlog": [
    {{
      "type": "Epic",
      "title": "Epic Title Here",
      "children": [
        {{
          "type": "Feature",
          "title": "Feature Title Here",
          "children": [
            {{
              "type": "User Story",
              "title": "User Story Title Here"
            }}
          ]
        }}
      ]
    }}
  ]
}}

--- HIGH-LEVEL PLAN ---
{plan}

--- ORIGINAL REQUIREMENTS TEXT ---
{original_text}
"""
    return call_generative_agent([{"text": prompt}], is_json=True)

# --- Helper functions (add_ids_to_backlog, Azure DevOps, etc.) ---
def add_ids_to_backlog(items):
    if not isinstance(items, list): return []
    for i, epic in enumerate(items):
        epic['id'] = f"E-{i+1}"
        for j, feature in enumerate(epic.get('children', []), 1):
            feature['id'] = f"F-{j}"
            for k, story in enumerate(feature.get('children', []), 1):
                story['id'] = f"US-{k}"
    return items

def get_ado_headers():
    pat_b64 = base64.b64encode(f":{ADO_PAT}".encode('utf-8')).decode('utf-8')
    return {'Content-Type': 'application/json-patch+json', 'Authorization': f'Basic {pat_b64}'}

def create_ado_work_item(item_type, title, parent_url=None):
    url = f"{ADO_ORGANIZATION_URL}/{ADO_PROJECT_NAME}/_apis/wit/workitems/${item_type}?api-version=7.1-preview.3"
    payload = [{"op": "add", "path": "/fields/System.Title", "value": title}]
    if parent_url:
        payload.append({"op": "add", "path": "/relations/-", "value": {"rel": "System.LinkTypes.Hierarchy-Reverse", "url": parent_url}})
    try:
        response = requests.post(url, headers=get_ado_headers(), data=json.dumps(payload))
        response.raise_for_status()
        return response.json()['url']
    except requests.exceptions.RequestException as e:
        print(f"ERROR creating ADO item: {e.response.text}")
        return None

def process_backlog_for_ado(backlog):
    print("Starting Azure DevOps work item creation...")
    for epic in backlog:
        epic_url = create_ado_work_item("Epic", epic['title'])
        if epic_url:
            for feature in epic.get('children', []):
                feature_url = create_ado_work_item("Feature", feature['title'], parent_url=epic_url)
                if feature_url:
                    for story in feature.get('children', []):
                        create_ado_work_item("User Story", story['title'], parent_url=feature_url)
    print("Finished Azure DevOps work item creation.")

def extract_mermaid_code(text):
    if not text:
        return ""
    match = re.search(r"```mermaid\n([\s\S]*?)```", text)
    if match:
        return match.group(1).strip()
    # Remove any standalone ```
    text = re.sub(r"```", "", text)
    return text.strip()

def render_mermaid_to_png(mermaid_code: str) -> BytesIO:
    url = "https://kroki.io/mermaid/png"
    headers = {"Content-Type": "text/plain"}
    response = requests.post(url, data=mermaid_code.encode("utf-8"), headers=headers)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        raise ValueError(f"Kroki diagram generation failed: {response.status_code} - {response.text}")

# --- API Endpoints ---

@app.route("/")
def index():
    return "<h1>AI Business Analyst Backend is running!</h1>"

@app.route("/api/generate", methods=['POST'])
def orchestrator():
    print("\n--- ORCHESTRATOR: New generation task started ---")
    if 'file' not in request.files:
        print("No file part in the request")
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    text_content, images, error = agent_extract_content(io.BytesIO(file.read()), file.filename)
    if error:
        print(f"Extraction error: {error}")
        return jsonify({"error": error}), 500

    # 1. Planning Agent
    plan, error = agent_planner(text_content, images)
    if error:
        print(f"Planning Agent failed: {error}")
        return jsonify({"error": f"Planning Agent failed: {error}"}), 500

    # 2. Specialist Agents
    trd, err_trd = agent_trd_writer(plan, text_content)
    hld, err_hld = agent_diagrammer(plan, "HLD")
    lld, err_lld = agent_diagrammer(plan, "LLD")
    backlog_json, err_backlog = agent_backlog_creator(plan, text_content)

    print(f"TRD: {trd}\nHLD: {hld}\nLLD: {lld}\nBacklog JSON: {backlog_json}")
    print(f"Errors: trd={err_trd}, hld={err_hld}, lld={err_lld}, backlog={err_backlog}")

    if any([err_trd, err_hld, err_lld, err_backlog]):
        print("One or more specialist agents failed.")
        return jsonify({"error": "One or more specialist agents failed."}), 500

    # 3. Final Assembly
    try:
        backlog_data = json.loads(backlog_json)
        actual_backlog_list = backlog_data.get('backlog', [])
        if not isinstance(actual_backlog_list, list):
            print("Warning: Backlog from AI is not a list, correcting.")
            actual_backlog_list = []

        hld_clean = extract_mermaid_code(hld)
        lld_clean = extract_mermaid_code(lld)
        print(f"Cleaned HLD: {hld_clean}\nCleaned LLD: {lld_clean}")

        final_response = {
            "trd": trd or "",
            "hld": hld_clean or "",
            "lld": lld_clean or "",
            "images": images,
            "backlog": add_ids_to_backlog(actual_backlog_list)
        }
    except json.JSONDecodeError as jde:
        print(f"Backlog Creator Agent returned invalid JSON: {backlog_json}\nError: {jde}")
        return jsonify({"error": "Backlog Creator Agent returned invalid JSON."}), 500
    except Exception as e:
        import traceback
        print(f"Error during final assembly: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Error during final assembly: {e}"}), 500

    print("--- ORCHESTRATOR: All agents completed. Task finished. ---\n")
    
    # Save analysis results
    analysis_id = save_analysis_results(final_response, text_content, file.filename if file else "requirements.txt")
    final_response["analysis_id"] = analysis_id
    
    return jsonify(final_response)

@app.route("/api/render_mermaid", methods=["POST"])
def render_mermaid():
    data = request.get_json()
    mermaid_code = data.get("code")
    if not mermaid_code:
        return jsonify({"error": "No Mermaid code provided"}), 400
    try:
        png_bytes = render_mermaid_to_png(mermaid_code)
        png_bytes.seek(0)
        return send_file(png_bytes, mimetype="image/png")
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/convert_to_docx", methods=['POST'])
def convert_to_docx():
    """Convert markdown content to DOCX format"""
    try:
        data = request.get_json()
        markdown_content = data.get('markdown', '')
        
        if not markdown_content:
            return jsonify({"error": "No markdown content provided"}), 400
        
        # Convert markdown to DOCX
        doc = markdown_to_docx(markdown_content)
        
        # Save to BytesIO
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return send_file(
            docx_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Technical_Requirements_Document.docx'
        )
        
    except Exception as e:
        print(f"Error converting to DOCX: {e}")
        return jsonify({"error": f"Failed to convert to DOCX: {str(e)}"}), 500

@app.route("/api/approve", methods=['POST'])
def send_for_approval():
    if not all([ACS_CONNECTION_STRING, ACS_SENDER_ADDRESS, APPROVAL_RECIPIENT_EMAIL]):
        return jsonify({"error": "Azure Communication Services is not configured on the server."}), 500
    
    request_data = request.get_json()
    documents = request_data.get('documents')
    if not documents:
        return jsonify({"error": "No documents provided for approval."}), 400

    approval_id = str(uuid.uuid4())
    approval_statuses[approval_id] = {"status": "pending", "data": documents}
    
    approve_url = f"{BACKEND_BASE_URL}/api/approval_response?id={approval_id}&decision=approved"
    reject_url = f"{BACKEND_BASE_URL}/api/approval_response?id={approval_id}&decision=rejected"

    email_content_html = f"""
    <html><body>
        <h2>Approval Request for Generated Business Artifacts</h2>
        <p>Please review the generated documents and approve or reject.</p>
        <p>
            <a href="{approve_url}" style="padding: 10px 15px; background-color: #28a745; color: white; text-decoration: none; border-radius: 5px;">Approve</a>
            <a href="{reject_url}" style="padding: 10px 15px; background-color: #dc3545; color: white; text-decoration: none; border-radius: 5px; margin-left: 10px;">Reject</a>
        </p>
        <hr><h3>Technical Requirements Document</h3><pre>{documents.get('trd', 'Not provided.')}</pre>
    </body></html>"""
    
    try:
        email_client = EmailClient.from_connection_string(ACS_CONNECTION_STRING)
        message = {
            "content": {"subject": "Approval Required: AI Generated Business Artifacts", "html": email_content_html},
            "recipients": {"to": [{"address": APPROVAL_RECIPIENT_EMAIL}]},
            "senderAddress": ACS_SENDER_ADDRESS
        }
        poller = email_client.begin_send(message)
        poller.result()
        print(f"Email sent for approval_id {approval_id}")
        return jsonify({"approval_id": approval_id, "status": "pending"})
    except Exception as e:
        print(f"Could not send email: {e}")
        return jsonify({"error": "Failed to send approval email."}), 500

@app.route("/api/approval_response")
def handle_approval_response():
    approval_id = request.args.get('id')
    decision = request.args.get('decision')

    if not approval_id or approval_id not in approval_statuses:
        return "<h1>Error: Invalid or expired approval request.</h1>", 400
    
    if approval_statuses[approval_id]["status"] != "pending":
         return "<h1>This request has already been processed.</h1>", 200

    approval_statuses[approval_id]["status"] = decision
    print(f"Approval ID {approval_id} has been {decision}.")

    if decision == 'approved':
        if not all([ADO_ORGANIZATION_URL, ADO_PROJECT_NAME, ADO_PAT]):
            print("ERROR: Azure DevOps environment variables are not set.")
            approval_statuses[approval_id]["status"] = "ado_failed"
        else:
            backlog_to_create = approval_statuses[approval_id]["data"]["backlog"]
            try:
                process_backlog_for_ado(backlog_to_create)
                approval_statuses[approval_id]["status"] = "approved_and_created"
            except Exception as e:
                print(f"ERROR: An exception occurred during ADO processing: {e}")
                approval_statuses[approval_id]["status"] = "ado_failed"

    confirmation_message = f"The request has been recorded as: <strong>{decision.capitalize()}</strong>."
    if approval_statuses[approval_id]["status"] == "approved_and_created":
        confirmation_message += "<p>Work items have been successfully created in Azure DevOps.</p>"
    elif approval_statuses[approval_id]["status"] == "ado_failed":
         confirmation_message += "<p style='color:red;'>However, there was an error creating the work items in Azure DevOps. Please check the server logs.</p>"

    return f"<html><body><h1>Thank you!</h1><p>{confirmation_message}</p><p>You can now close this window.</p></body></html>", 200

@app.route("/api/approval_status/<approval_id>")
def get_approval_status(approval_id):
    status_info = approval_statuses.get(approval_id)
    if not status_info:
        return jsonify({"error": "Approval ID not found."}), 404
    return jsonify({"status": status_info["status"]})

@app.route("/api/search", methods=['POST'])
def semantic_search():
    """Perform semantic search across documents and analyses"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        collection = data.get('collection', 'documents')
        n_results = data.get('n_results', 5)
        
        if not query:
            return jsonify({"error": "Query is required"}), 400
        
        # Perform vector search
        results = search_vector_db(query, collection, n_results)
        
        return jsonify({
            "query": query,
            "results": results,
            "total_results": len(results.get('documents', []))
        })
        
    except Exception as e:
        print(f"Error in semantic search: {e}")
        return jsonify({"error": f"Search failed: {str(e)}"}), 500

@app.route("/api/vector/collections", methods=['GET'])
def get_collections():
    """Get all Qdrant collections"""
    try:
        collections = qdrant_client.get_collections()
        return jsonify({
            "collections": [col.name for col in collections.collections]
        })
    except Exception as e:
        print(f"Error getting collections: {e}")
        return jsonify({"error": f"Failed to get collections: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/info", methods=['GET'])
def get_collection_info(collection_name):
    """Get information about a specific collection"""
    try:
        collection_info = qdrant_client.get_collection(collection_name)
        return jsonify({
            "name": collection_name,
            "vectors_count": collection_info.vectors_count,
            "points_count": collection_info.points_count,
            "status": collection_info.status
        })
    except Exception as e:
        print(f"Error getting collection info: {e}")
        return jsonify({"error": f"Failed to get collection info: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points", methods=['GET'])
def get_collection_points(collection_name):
    """Get points from a collection"""
    try:
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        points = qdrant_client.scroll(
            collection_name=collection_name,
            limit=limit,
            offset=offset
        )
        
        return jsonify({
            "collection": collection_name,
            "points": [
                {
                    "id": point.id,
                    "payload": point.payload,
                    "score": getattr(point, 'score', None)
                }
                for point in points[0]
            ],
            "total": len(points[0])
        })
    except Exception as e:
        print(f"Error getting collection points: {e}")
        return jsonify({"error": f"Failed to get collection points: {str(e)}"}), 500

@app.route("/api/vector/collections/<collection_name>/points/<point_id>", methods=['DELETE'])
def delete_point(collection_name, point_id):
    """Delete a point from a collection"""
    try:
        success = delete_from_vector_db(point_id, collection_name)
        if success:
            return jsonify({"message": f"Point {point_id} deleted successfully"})
        else:
            return jsonify({"error": "Failed to delete point"}), 500
    except Exception as e:
        print(f"Error deleting point: {e}")
        return jsonify({"error": f"Failed to delete point: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
